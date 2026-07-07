"""
Generate cost budget report (DOCX + PDF) for leadership review.

Source: D:\\code\\labor-survey-ai\\reports\\cost-budget-20260622.md
Outputs:
  - D:\\code\\labor-survey-ai\\reports\\成本预算-三档用量-20260622.docx
  - D:\\code\\labor-survey-ai\\reports\\成本预算-三档用量-20260622.pdf
"""

from __future__ import annotations

import os
from copy import deepcopy
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm, mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

OUTPUT_DIR = r"D:\code\labor-survey-ai\reports"
OUTPUT_BASENAME = "成本预算-三档用量-20260622"
OUTPUT_DOCX = os.path.join(OUTPUT_DIR, f"{OUTPUT_BASENAME}.docx")
OUTPUT_PDF = os.path.join(OUTPUT_DIR, f"{OUTPUT_BASENAME}.pdf")

REPORT_TITLE = "劳动力调查 AI 助手 · 成本预算报告"
REPORT_SUBTITLE = "三档用量 × 三档人数对照"
ORG_NAME = "国家统计局贵阳调查队"
AUTHOR = "刘佳旭"
DATE = "2026-06-22"
FILE_LABEL = "成本预算-三档用量-20260622"

GOV_BLUE = RGBColor(0x1E, 0x3A, 0x5F)
GOV_BLUE_HEX = "#1E3A5F"
GOV_BLUE_LIGHT_HEX = "#E8EEF5"
GOV_BLUE_HEADER_HEX = "#1E3A5F"
GOV_BLUE_TEXT_HEX = "#FFFFFF"
GREEN_HIGHLIGHT_HEX = "#C6EFCE"
GREEN_TEXT_HEX = "#006100"
YELLOW_HIGHLIGHT_HEX = "#FFEB9C"
YELLOW_TEXT_HEX = "#9C5700"


# ---------------------------------------------------------------------------
# Common data (the "9-grid" matrix used in both DOCX and PDF)
# ---------------------------------------------------------------------------

GRID_DATA: list[list[str]] = [
    ["人数 \\ 月量", "1000 次", "7000 次", "20万次"],
    ["100 人（处室）", "¥3", "¥24", "¥698 ⚠"],
    ["700 人（市级）", "¥66", "¥87 ⭐", "¥761"],
    ["2 万人（全省/全市）", "¥371", "¥392", "¥1066"],
]
# (row_idx, col_idx) zero-based
GRID_HIGHLIGHTS: list[tuple[int, int, str, str]] = [
    (1, 3, YELLOW_HIGHLIGHT_HEX, YELLOW_TEXT_HEX),   # 100人 × 20万次
    (2, 2, GREEN_HIGHLIGHT_HEX, GREEN_TEXT_HEX),     # 700人 × 7000次
]


# ---------------------------------------------------------------------------
# DOCX generation
# ---------------------------------------------------------------------------

def _set_run_font(run, name_cn: str, name_en: str, size_pt: float, color: RGBColor | None = None, bold: bool = False):
    run.font.name = name_en
    run.font.size = Pt(size_pt)
    if bold:
        run.bold = True
    if color is not None:
        run.font.color.rgb = color
    # CJK font binding
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name_cn)
    rFonts.set(qn("w:ascii"), name_en)
    rFonts.set(qn("w:hAnsi"), name_en)


def _add_paragraph(doc, text="", style=None):
    if style is not None:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
    return p


def _shade_cell(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex.lstrip("#"))
    tcPr.append(shd)


def _set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Border in eighths-of-a-point. None = no change. False = remove."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for name, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        if val is None:
            continue
        existing = tcBorders.find(qn(f"w:{name}"))
        if existing is not None:
            tcBorders.remove(existing)
        b = OxmlElement(f"w:{name}")
        if val is False:
            b.set(qn("w:val"), "nil")
        else:
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), str(val))
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), GOV_BLUE_HEX.lstrip("#"))
        tcBorders.append(b)


def _set_table_three_line(table, header_rows: int = 1):
    """三线表：顶线 + 表头底线 + 表底线，无竖线。"""
    n_rows = len(table.rows)
    n_cols = len(table.columns)
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            top = None
            bottom = None
            if r_idx == 0:
                top = 12  # 1.5pt
            if r_idx == header_rows - 1:
                bottom = 8  # 1pt
            if r_idx == n_rows - 1:
                bottom = 12
            _set_cell_borders(cell, top=top, bottom=bottom, left=False, right=False)
            # remove vertical default border too
    # Inner horizontal borders: clear all
    tbl = table._tbl
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        tblBorders.append(b)
    existing = tbl.find(qn("w:tblPr"))
    if existing is not None:
        old = existing.find(qn("w:tblBorders"))
        if old is not None:
            existing.remove(old)
        existing.append(tblBorders)


def _fill_cell(cell, text: str, *, bold: bool = False, color: RGBColor | None = None,
               align=WD_ALIGN_PARAGRAPH.CENTER, size: float = 10, font_cn: str = "微软雅黑"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    _set_run_font(run, font_cn, font_cn, size, color=color, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _setup_doc_styles(doc):
    # Default style tweak
    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal.font.size = Pt(11)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "微软雅黑")
    rFonts.set(qn("w:ascii"), "微软雅黑")
    rFonts.set(qn("w:hAnsi"), "微软雅黑")


def _add_page_number_field(paragraph, prefix_text: str = "第 ", suffix_text: str = " 页 / 共 "):
    run = paragraph.add_run(prefix_text)
    _set_run_font(run, "微软雅黑", "微软雅黑", 9, color=GOV_BLUE)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    page_run = paragraph.add_run()
    page_run._r.append(fld_begin)
    page_run._r.append(instr)
    page_run._r.append(fld_sep)
    page_run._r.append(fld_end)
    _set_run_font(page_run, "微软雅黑", "微软雅黑", 9, color=GOV_BLUE)

    run2 = paragraph.add_run(suffix_text)
    _set_run_font(run2, "微软雅黑", "微软雅黑", 9, color=GOV_BLUE)

    fld_begin2 = OxmlElement("w:fldChar")
    fld_begin2.set(qn("w:fldCharType"), "begin")
    instr2 = OxmlElement("w:instrText")
    instr2.set(qn("xml:space"), "preserve")
    instr2.text = " NUMPAGES "
    fld_sep2 = OxmlElement("w:fldChar")
    fld_sep2.set(qn("w:fldCharType"), "separate")
    fld_end2 = OxmlElement("w:fldChar")
    fld_end2.set(qn("w:fldCharType"), "end")
    page_run2 = paragraph.add_run()
    page_run2._r.append(fld_begin2)
    page_run2._r.append(instr2)
    page_run2._r.append(fld_sep2)
    page_run2._r.append(fld_end2)
    _set_run_font(page_run2, "微软雅黑", "微软雅黑", 9, color=GOV_BLUE)

    run3 = paragraph.add_run(" 页")
    _set_run_font(run3, "微软雅黑", "微软雅黑", 9, color=GOV_BLUE)


def _setup_header_footer(doc):
    for section in doc.sections:
        # 页眉：报告标题（居左）+ 页码（居右）
        header = section.header
        # Header is a single paragraph; use tab stops to push page-number to right
        hp = header.paragraphs[0]
        hp.text = ""
        tab_stops = hp.paragraph_format.tab_stops
        # Use a right tab at the right margin
        tab_stops.add_tab_stop(Cm(16), WD_ALIGN_PARAGRAPH.RIGHT)
        run = hp.add_run(REPORT_TITLE)
        _set_run_font(run, "微软雅黑", "微软雅黑", 9, color=GOV_BLUE)
        run2 = hp.add_run("\t")
        _set_run_font(run2, "微软雅黑", "微软雅黑", 9, color=GOV_BLUE)
        _add_page_number_field(hp, prefix_text="")

        # 底部加一条蓝色横线效果靠 paragraph border
        pPr = hp._p.get_or_add_pPr()
        pBdr = pPr.find(qn("w:pBdr"))
        if pBdr is None:
            pBdr = OxmlElement("w:pBdr")
            pPr.append(pBdr)
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), GOV_BLUE_HEX.lstrip("#"))
        pBdr.append(bottom)

        # 页脚：编制人 + 日期 + 文件名
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run(f"编制人：{AUTHOR}　|　编制日期：{DATE}　|　文件名：{FILE_LABEL}")
        _set_run_font(run, "微软雅黑", "微软雅黑", 9, color=GOV_BLUE)


def _add_cover_page(doc):
    # 顶部留白
    for _ in range(4):
        doc.add_paragraph()

    # 单位
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(ORG_NAME)
    _set_run_font(run, "微软雅黑", "微软雅黑", 14, color=GOV_BLUE, bold=True)

    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(48)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(REPORT_TITLE)
    _set_run_font(run, "宋体", "Times New Roman", 28, color=GOV_BLUE, bold=True)

    # 副标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(48)
    run = p.add_run(REPORT_SUBTITLE)
    _set_run_font(run, "微软雅黑", "微软雅黑", 16, color=GOV_BLUE)

    # 装饰线
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), GOV_BLUE_HEX.lstrip("#"))
    pBdr.append(bottom)
    pPr.append(pBdr)
    run = p.add_run("　" * 20)
    _set_run_font(run, "微软雅黑", "微软雅黑", 12, color=GOV_BLUE)

    # 留白
    for _ in range(8):
        doc.add_paragraph()

    # 编制人 / 日期 信息
    info_table = doc.add_table(rows=2, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False
    widths = [Cm(4), Cm(6)]
    for row in info_table.rows:
        for i, c in enumerate(row.cells):
            c.width = widths[i]
    # 清除所有边框
    tbl = info_table._tbl
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        tblBorders.append(b)
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)

    _fill_cell(info_table.rows[0].cells[0], "编 制 人", bold=True, color=GOV_BLUE, size=12)
    _fill_cell(info_table.rows[0].cells[1], AUTHOR, color=GOV_BLUE, size=12)
    _fill_cell(info_table.rows[1].cells[0], "编制日期", bold=True, color=GOV_BLUE, size=12)
    _fill_cell(info_table.rows[1].cells[1], DATE, color=GOV_BLUE, size=12)

    # 末尾留白 + 强制分页
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _add_h1(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    _set_run_font(run, "宋体", "Times New Roman", 16, color=GOV_BLUE, bold=True)
    # 左侧蓝色竖条
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), GOV_BLUE_HEX.lstrip("#"))
    pBdr.append(left)
    return p


def _add_h2(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_run_font(run, "宋体", "Times New Roman", 13, color=GOV_BLUE, bold=True)
    return p


def _add_body(doc, text: str, *, indent: bool = True, bold: bool = False, color: RGBColor | None = None, size: float = 11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.first_line_indent = Pt(22)
    run = p.add_run(text)
    _set_run_font(run, "微软雅黑", "微软雅黑", size, color=color, bold=bold)
    return p


def _add_bullet(doc, text: str, *, size: float = 11):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    _set_run_font(run, "微软雅黑", "微软雅黑", size)
    return p


def _add_table(doc, data: list[list[str]], *, header: bool = True, highlights: list[tuple[int, int, str, str]] | None = None,
               col_widths: list[float] | None = None, font_size: float = 10):
    rows = len(data)
    cols = len(data[0])
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if col_widths is not None:
        for r in table.rows:
            for i, c in enumerate(r.cells):
                c.width = Cm(col_widths[i])
    # 清除默认边框，再设三线表
    _set_table_three_line(table, header_rows=1 if header else 0)
    # 填充内容
    for r_idx, row in enumerate(data):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            is_header = header and r_idx == 0
            if is_header:
                _fill_cell(cell, val, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=font_size, font_cn="微软雅黑")
                _shade_cell(cell, GOV_BLUE_HEX.lstrip("#"))
            else:
                _fill_cell(cell, val, size=font_size)
    # 高亮
    if highlights:
        for r_idx, c_idx, fill, font_color in highlights:
            cell = table.rows[r_idx].cells[c_idx]
            _shade_cell(cell, fill)
            # 改字体颜色
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor.from_string(font_color.lstrip("#"))
                    run.bold = True
    return table


def _page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build_docx() -> str:
    doc = Document()
    _setup_doc_styles(doc)

    # 页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.4)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    _setup_header_footer(doc)

    # 封面
    _add_cover_page(doc)

    # 一、摘要
    _add_h1(doc, "一、摘要")
    _add_body(doc, "本报告基于已抓取的 DeepSeek v4-flash 实时价格（api-docs.deepseek.com）以及训练数据估算，"
                   "测算劳动力调查 AI 助手在三档人数（100人/700人/2万人）× 三档月用量（1000/7000/20万次）共 9 档组合下的"
                   "月度成本。", indent=True)
    _add_body(doc, "采购建议档：", indent=True, bold=True)
    _add_body(doc, "700 人（市级调查员）× 7000 次/月 ≈ ¥87/月（API ¥24 + 云资源 ¥63）", indent=False, bold=True, color=GOV_BLUE)
    _add_body(doc, "该档覆盖了「市级调查员每月一次集中查询」这一核心真实场景，且 100 人档可直接降级为本地电脑 + Cloudflare Tunnel，"
                   "无任何云资源开支。", indent=True)

    _add_h2(doc, "1.1 关键风险提示")
    risks = [
        ("数据合规", "本系统架构已严格隔离居民个人信息（H5 → 后端 API → 阿里云百炼 / DeepSeek，仅传查询文本）；"
                     "扩量至市级或省级时，需在采购合同中明确数据流向条款。"),
        ("政府采购流程", "云资源月租达一定金额通常需走政府采购或协议供货，域名备案需 15-20 工作日，"
                          "预算应包含备案申请工时。API Key 必须开在单位名下，避免离职交接问题。"),
        ("量纲校验", "100 人 × 20 万次/月（每人 67 次/天）属异常档，仅作 API 容量上限测试；"
                    "2 万 × 1000 次/月（每 20 人 1 次/月）使用率过低，建议合并到 7000 次档。"),
    ]
    for title, body in risks:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(22)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"【{title}】")
        _set_run_font(run, "微软雅黑", "微软雅黑", 11, color=GOV_BLUE, bold=True)
        run2 = p.add_run(body)
        _set_run_font(run2, "微软雅黑", "微软雅黑", 11)

    # 二、9 档对照汇总表（重点页）
    _add_h1(doc, "二、9 档对照汇总表")
    _add_body(doc, "绿色高亮为采购建议档，黄色高亮为异常档（仅作压力测试参考）。", indent=True)
    table = _add_table(doc, GRID_DATA, col_widths=[4.5, 3.5, 3.5, 3.5], highlights=GRID_HIGHLIGHTS)
    # Make the grid table larger font
    for row in table.rows:
        for c in row.cells:
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("采购建议档：700 人 × 7000 次/月 ≈ ¥87/月（API ¥24 + 云资源 ¥63）")
    _set_run_font(run, "微软雅黑", "微软雅黑", 11, color=RGBColor.from_string(GREEN_TEXT_HEX.lstrip("#")), bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(22)
    run = p.add_run("⚠ 异常档：100 人 × 20 万次 = ¥698/月，仅作 API 容量上限参考，不构成实际采购建议。")
    _set_run_font(run, "微软雅黑", "微软雅黑", 11, color=RGBColor.from_string(YELLOW_TEXT_HEX.lstrip("#")), bold=True)

    # 三、价格弹性
    _add_h2(doc, "2.1 价格弹性分析")
    _add_bullet(doc, "API 费用（DeepSeek + DashScope）随用量线性增长，几乎纯按量计费。")
    _add_bullet(doc, "云资源费用为阶梯式：100 人档 ¥0（本地电脑）；700 人档 ~¥62.5/月（ECS + 域名）；2 万档 ~¥367/月（ECS + SLB + RDS + CDN + 域名）。")
    _add_bullet(doc, "采购建议档的 ¥87/月：API ¥24（占 28%）+ 云资源 ¥63（占 72%）。降低 API 用量对总成本影响有限，应聚焦云资源选型。")

    # 三、详细分档说明
    _add_h1(doc, "三、详细分档说明")

    # 3.1 档位 A
    _add_h2(doc, "3.1 档位 A：100 人（处室自用）")
    _add_body(doc, "部署：本地电脑 + Cloudflare Tunnel quick（与现状一致）。", indent=True)
    _add_body(doc, "不需要：ECS、域名、备案、CDN。", indent=True)
    _add_h2(doc, "费用明细（单位：元/月）")
    _add_table(doc, [
        ["开销项目", "1000 次", "7000 次", "20万次 ⚠"],
        ["LLM 输入（DeepSeek v4-flash）", "¥1.70", "¥11.90", "¥340.00"],
        ["LLM 输出（DeepSeek v4-flash）", "¥0.60", "¥4.20", "¥120.00"],
        ["Embedding（DashScope v3）", "¥1.19", "¥8.33", "¥238.00"],
        ["Cloudflare Tunnel", "免费", "免费", "免费"],
        ["合计", "≈ ¥3/月", "≈ ¥24/月", "≈ ¥698/月"],
    ], col_widths=[6.5, 3.0, 3.0, 3.0])
    _add_body(doc, "⚠ 100 人 × 20 万次 = 每人 2000 次/月（≈67 次/天），属明显异常场景，仅作 API 容量上限参考。", indent=True, color=RGBColor.from_string(YELLOW_TEXT_HEX.lstrip("#")))

    _add_h2(doc, "适用判定")
    _add_bullet(doc, "实际处室用量（每月 100-700 次）：API 月费 < ¥3，可忽略不计。")
    _add_bullet(doc, "异常档 ¥698：API 费用已超过部署 ECS 的成本，此规模应迁云。")

    # 3.2 档位 B
    _add_h2(doc, "3.2 档位 B：700 人（市级使用）")
    _add_body(doc, "部署：阿里云 ECS（2核4G 突发 t6）+ Cloudflare Tunnel named + 自有域名。", indent=True)
    _add_body(doc, "域名备案：政府单位走阿里云 ICP 备案系统（约 15-20 工作日，免费）。", indent=True)
    _add_body(doc, "HTTPS：Let's Encrypt 自动续期。", indent=True)
    _add_body(doc, "不需要：SLB（单机够用）、CDN（H5 体积小）。", indent=True)
    _add_h2(doc, "费用明细（单位：元/月）")
    _add_table(doc, [
        ["开销项目", "1000 次", "7000 次", "20万次"],
        ["LLM 输入", "¥1.70", "¥11.90", "¥340.00"],
        ["LLM 输出", "¥0.60", "¥4.20", "¥120.00"],
        ["Embedding", "¥1.19", "¥8.33", "¥238.00"],
        ["ECS（2核4G t6）", "¥60", "¥60", "¥60"],
        [".cn 域名（年付折月）", "¥2.5", "¥2.5", "¥2.5"],
        ["ICP 备案", "一次性免费", "—", "—"],
        ["Cloudflare Tunnel named", "免费", "免费", "免费"],
        ["SSL 证书", "免费", "免费", "免费"],
        ["合计", "≈ ¥66/月", "≈ ¥87/月", "≈ ¥761/月"],
    ], col_widths=[6.5, 3.0, 3.0, 3.0])

    _add_h2(doc, "适用判定")
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("700 × 7000 次/月：月费 ≈ ¥87，")
    _set_run_font(run, "微软雅黑", "微软雅黑", 11)
    run = p.add_run("推荐采购档")
    _set_run_font(run, "微软雅黑", "微软雅黑", 11, color=RGBColor.from_string(GREEN_TEXT_HEX.lstrip("#")), bold=True)
    _add_bullet(doc, "700 × 20 万次/月：月费 ≈ ¥761，需要提前扩容 ECS 至 4 核。")
    _add_body(doc, "高并发期（7 天承担 60-70% 流量）峰值约 200 次/天（20 万档）。单 ECS 2核4G 跑 DeepSeek 转发绰绰有余，"
                   "瓶颈在 DeepSeek 端 QPS 限额，不在本机。", indent=True)

    # 3.3 档位 C
    _add_h2(doc, "3.3 档位 C：2 万人（全省/全市调查员）")
    _add_body(doc, "部署：阿里云 ECS（4核8G 计算型 c7）+ SLB + Cloudflare Tunnel named + 域名 + 备案。", indent=True)
    _add_body(doc, "CDN：H5 静态资源走阿里云 CDN（按流量计费）。", indent=True)
    _add_body(doc, "向量库升级：Chroma → 独立服务（Milvus 或 Zvec 备选方案）。", indent=True)
    _add_body(doc, "数据库：用户反馈数据上云数据库 RDS（按量计费）。", indent=True)
    _add_h2(doc, "费用明细（单位：元/月）")
    _add_table(doc, [
        ["开销项目", "1000 次", "7000 次", "20万次"],
        ["LLM 输入", "¥1.70", "¥11.90", "¥340.00"],
        ["LLM 输出", "¥0.60", "¥4.20", "¥120.00"],
        ["Embedding", "¥1.19", "¥8.33", "¥238.00"],
        ["ECS（4核8G c7）", "¥300", "¥300", "¥300"],
        ["SLB（含流量）", "¥30", "¥30", "¥30"],
        [".cn 域名", "¥2.5", "¥2.5", "¥2.5"],
        ["CDN（H5 静态）", "~¥5", "~¥5", "~¥5"],
        ["RDS MySQL（反馈数据）", "¥30", "¥30", "¥30"],
        ["ICP 备案", "一次性免费", "—", "—"],
        ["Cloudflare Tunnel named", "免费", "免费", "免费"],
        ["合计", "≈ ¥371/月", "≈ ¥392/月", "≈ ¥1066/月"],
    ], col_widths=[6.5, 3.0, 3.0, 3.0])

    _add_h2(doc, "适用判定")
    _add_bullet(doc, "2 万 × 7000 次/月（每人每月 0.35 次 ≈ 每 3 月 1 次）：月费 ≈ ¥392。")
    _add_bullet(doc, "2 万 × 20 万次/月（每人每月 10 次 ≈ 每周 2 次）：月费 ≈ ¥1066。")
    _add_body(doc, "⚠ 高并发期 7 天承担 14 万次 ≈ 每天 2 万次 ≈ 每分钟 35 次（按 8 小时工作日计）。"
                   "4核8G 单机压力测试需做，极端情况可能需要双 ECS + SLB 横向扩展（月费翻倍至 ¥1500+）。",
              indent=True, color=RGBColor.from_string(YELLOW_TEXT_HEX.lstrip("#")))

    # 四、附录 A：价格核对清单
    _add_h1(doc, "四、附录 A：价格核对清单")
    _add_body(doc, "以下 7 项服务因登录墙 / JS 渲染，自动脚本无法获取精确价格。打印本表后请人工登录核对。", indent=True)
    _add_table(doc, [
        ["#", "服务", "核对路径", "期望价格", "实际价格（请填）"],
        ["1", "DashScope text-embedding-v3", "https://dashscope.console.aliyun.com/billing", "¥0.7/M tokens", "　　待填"],
        ["2", "阿里云 ECS t6（2核4G）", "https://www.aliyun.com/product/ecs", "~¥60/月", "　　待填"],
        ["3", "阿里云 ECS c7（4核8G）", "同上 → 计算型", "~¥300/月", "　　待填"],
        ["4", "阿里云 SLB", "https://www.aliyun.com/product/slb", "~¥20/月 + 流量", "　　待填"],
        ["5", ".cn 域名", "https://wanwang.aliyun.com", "~¥30/年", "　　待填"],
        ["6", "阿里云 CDN", "https://www.aliyun.com/product/cdn", "~¥0.24/GB", "　　待填"],
        ["7", "阿里云 RDS MySQL", "https://www.aliyun.com/product/rds/mysql", "~¥30/月起", "　　待填"],
    ], col_widths=[0.8, 3.6, 5.5, 2.6, 2.5])
    _add_body(doc, "重算公式（每千次）：", indent=True, bold=True)
    _add_body(doc, "LLM_每千次 = (1.7 × LLM_输入单价 + 0.3 × LLM_输出单价)", indent=False)
    _add_body(doc, "Embedding_每千次 = 1.7 × Embedding_单价", indent=False)
    _add_body(doc, "总月费 = (月次数 ÷ 1000) × (LLM_每千次 + Embedding_每千次) + 云资源固定费", indent=False)

    # 五、附录 B：已确认价格
    _add_h1(doc, "五、附录 B：已确认价格")
    _add_h2(doc, "5.1 DeepSeek v4-flash（自动抓取）")
    _add_table(doc, [
        ["项目", "美元原价", "换算 CNY（按 USD/CNY ≈ 7.2）"],
        ["输入（cache miss）", "$0.14 / M tokens", "≈ ¥1.0 / M tokens"],
        ["输出", "$0.28 / M tokens", "≈ ¥2.0 / M tokens"],
    ], col_widths=[6.0, 5.0, 4.5])
    _add_body(doc, "数据来源：https://api-docs.deepseek.com/quick_start/pricing （2026-06-22 抓取）", indent=True)
    _add_h2(doc, "5.2 Cloudflare Tunnel（公开文档）")
    _add_table(doc, [
        ["模式", "定价", "说明"],
        ["quick tunnel（一次性）", "免费", "无需账号，URL 随机"],
        ["named tunnel（长期）", "免费", "需 Cloudflare 账号 + 域名挂载"],
    ], col_widths=[5.0, 3.0, 7.5])
    _add_body(doc, "数据来源：https://developers.cloudflare.com/cloudflare-one/ （2026-06-22 抓取）", indent=True)
    _add_body(doc, "Let's Encrypt SSL 证书：免费（自动续期）。", indent=True)

    # 报告结尾
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("—— 报告结束 ——")
    _set_run_font(run, "微软雅黑", "微软雅黑", 12, color=GOV_BLUE, bold=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("本表可作为采购可行性报告附件提交。")
    _set_run_font(run2, "微软雅黑", "微软雅黑", 10, color=GOV_BLUE)

    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


# ---------------------------------------------------------------------------
# PDF generation
# ---------------------------------------------------------------------------

def _register_cn_font():
    """Register a CJK font. Use STSong-Light (built-in CID font)."""
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    return "STSong-Light"


def _pdf_para(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(text, style)


def _wrap_long_cell(text: str, style: ParagraphStyle) -> Paragraph:
    """Wrap a string as a Paragraph.

    If the text already contains a balanced <...> tag, trust the caller and
    only escape ampersands. Otherwise treat as plain text and escape HTML.
    """
    # Heuristic: if it looks like a Paragraph fragment (has tags), escape only &
    if "<" in text and ">" in text:
        safe = text.replace("&", "&amp;")
    else:
        safe = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    safe = safe.replace("\n", "<br/>")
    return Paragraph(safe, style)


_TOTAL_PAGES_HOLDER = [0]  # mutable container for two-pass page count


def _add_page_decorator(canvas, doc):
    canvas.saveState()
    width, height = A4

    # 页眉：报告标题（左）+ 页码（右）
    canvas.setFont("STSong-Light", 9)
    canvas.setFillColor(colors.HexColor(GOV_BLUE_HEX))
    canvas.drawString(2.0 * cm, height - 1.3 * cm, REPORT_TITLE)
    if _TOTAL_PAGES_HOLDER[0] > 0:
        page_text = f"第 {doc.page} 页 / 共 {_TOTAL_PAGES_HOLDER[0]} 页"
    else:
        page_text = f"第 {doc.page} 页"
    canvas.drawRightString(width - 2.0 * cm, height - 1.3 * cm, page_text)
    # header line
    canvas.setStrokeColor(colors.HexColor(GOV_BLUE_HEX))
    canvas.setLineWidth(0.6)
    canvas.line(2.0 * cm, height - 1.45 * cm, width - 2.0 * cm, height - 1.45 * cm)

    # 页脚
    canvas.setFont("STSong-Light", 9)
    canvas.setFillColor(colors.HexColor(GOV_BLUE_HEX))
    footer_text = f"编制人：{AUTHOR}　|　编制日期：{DATE}　|　文件名：{FILE_LABEL}"
    canvas.drawCentredString(width / 2.0, 1.3 * cm, footer_text)
    canvas.setStrokeColor(colors.HexColor(GOV_BLUE_HEX))
    canvas.setLineWidth(0.4)
    canvas.line(2.0 * cm, 1.5 * cm, width - 2.0 * cm, 1.5 * cm)
    canvas.restoreState()


def _styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    normal = ParagraphStyle(
        "CNBody",
        parent=base["BodyText"],
        fontName="STSong-Light",
        fontSize=11,
        leading=18,
        textColor=colors.black,
        alignment=0,  # left
        firstLineIndent=22,
    )
    h1 = ParagraphStyle(
        "CNH1",
        parent=base["Heading1"],
        fontName="STSong-Light",
        fontSize=16,
        leading=22,
        textColor=colors.HexColor(GOV_BLUE_HEX),
        spaceBefore=14,
        spaceAfter=8,
    )
    h2 = ParagraphStyle(
        "CNH2",
        parent=base["Heading2"],
        fontName="STSong-Light",
        fontSize=13,
        leading=18,
        textColor=colors.HexColor(GOV_BLUE_HEX),
        spaceBefore=10,
        spaceAfter=4,
    )
    cover_title = ParagraphStyle(
        "CoverTitle",
        fontName="STSong-Light",
        fontSize=28,
        leading=40,
        textColor=colors.HexColor(GOV_BLUE_HEX),
        alignment=1,
        spaceBefore=24,
        spaceAfter=12,
    )
    cover_sub = ParagraphStyle(
        "CoverSub",
        fontName="STSong-Light",
        fontSize=16,
        leading=24,
        textColor=colors.HexColor(GOV_BLUE_HEX),
        alignment=1,
        spaceAfter=24,
    )
    cover_org = ParagraphStyle(
        "CoverOrg",
        fontName="STSong-Light",
        fontSize=14,
        leading=20,
        textColor=colors.HexColor(GOV_BLUE_HEX),
        alignment=1,
        spaceAfter=8,
    )
    cover_meta = ParagraphStyle(
        "CoverMeta",
        fontName="STSong-Light",
        fontSize=12,
        leading=20,
        textColor=colors.HexColor(GOV_BLUE_HEX),
        alignment=1,
    )
    body_no_indent = ParagraphStyle(
        "CNBodyNoIndent",
        parent=normal,
        firstLineIndent=0,
    )
    body_bold = ParagraphStyle(
        "CNBodyBold",
        parent=normal,
        fontName="STSong-Light",
    )
    bullet = ParagraphStyle(
        "CNBullet",
        parent=body_no_indent,
        leftIndent=18,
        bulletIndent=6,
        spaceAfter=2,
    )
    cell_header = ParagraphStyle(
        "CellHeader",
        fontName="STSong-Light",
        fontSize=10,
        leading=14,
        alignment=1,
        textColor=colors.white,
    )
    cell_body = ParagraphStyle(
        "CellBody",
        fontName="STSong-Light",
        fontSize=10,
        leading=14,
        alignment=1,
    )
    cell_left = ParagraphStyle(
        "CellLeft",
        fontName="STSong-Light",
        fontSize=10,
        leading=14,
        alignment=0,
    )
    note_green = ParagraphStyle(
        "NoteGreen",
        fontName="STSong-Light",
        fontSize=11,
        leading=18,
        textColor=colors.HexColor(GREEN_TEXT_HEX),
        firstLineIndent=22,
    )
    note_yellow = ParagraphStyle(
        "NoteYellow",
        fontName="STSong-Light",
        fontSize=11,
        leading=18,
        textColor=colors.HexColor(YELLOW_TEXT_HEX),
        firstLineIndent=22,
    )
    return {
        "normal": normal,
        "no_indent": body_no_indent,
        "bold": body_bold,
        "bullet": bullet,
        "h1": h1,
        "h2": h2,
        "cover_title": cover_title,
        "cover_sub": cover_sub,
        "cover_org": cover_org,
        "cover_meta": cover_meta,
        "cell_header": cell_header,
        "cell_body": cell_body,
        "cell_left": cell_left,
        "note_green": note_green,
        "note_yellow": note_yellow,
    }


def _table_style_3line(header_rows: int = 1, body_rows: int = 0, col_widths: list[float] | None = None) -> TableStyle:
    """三线表：顶线 + 表头底线 + 表底线，无竖线。"""
    style_cmds = [
        ("FONT", (0, 0), (-1, -1), "STSong-Light", 10),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        # Header
        ("BACKGROUND", (0, 0), (-1, header_rows - 1), colors.HexColor(GOV_BLUE_HEX)),
        ("TEXTCOLOR", (0, 0), (-1, header_rows - 1), colors.white),
        # 顶线
        ("LINEABOVE", (0, 0), (-1, 0), 1.5, colors.HexColor(GOV_BLUE_HEX)),
        # 表头底线
        ("LINEBELOW", (0, header_rows - 1), (-1, header_rows - 1), 1.0, colors.HexColor(GOV_BLUE_HEX)),
        # 表底线
        ("LINEBELOW", (0, body_rows), (-1, body_rows), 1.5, colors.HexColor(GOV_BLUE_HEX)),
        # 移除竖线
        ("LINEBEFORE", (0, 0), (-1, -1), 0, colors.white),
        ("LINEAFTER", (0, 0), (-1, -1), 0, colors.white),
    ]
    return TableStyle(style_cmds)


def _make_table(data: list[list[str]], col_widths: list[float] | None, styles: dict, header_rows: int = 1,
                cell_styles: list[ParagraphStyle] | None = None) -> Table:
    """Wrap strings as Paragraphs so they flow nicely."""
    if cell_styles is None:
        cell_styles = [styles["cell_header"]] + [styles["cell_body"]] * (len(data) - 1)
    wrapped: list[list[Paragraph]] = []
    for r_idx, row in enumerate(data):
        style_for_row = cell_styles[0] if r_idx < header_rows else cell_styles[1]
        wrapped_row = []
        for c_idx, val in enumerate(row):
            if c_idx == 0 and r_idx >= header_rows and isinstance(val, str) and len(val) > 8:
                style = styles["cell_left"]
            else:
                style = style_for_row
            wrapped_row.append(_wrap_long_cell(val, style))
        wrapped.append(wrapped_row)
    t = Table(wrapped, colWidths=col_widths, repeatRows=header_rows)
    t.setStyle(_table_style_3line(header_rows=header_rows, body_rows=len(data) - 1))
    return t


def _build_story(styles: dict) -> list[Any]:
    """Return the platypus story (cover, body, appendices, end)."""
    story: list[Any] = []
    # 封面
    story.append(Spacer(1, 3 * cm))
    story.append(_pdf_para(ORG_NAME, styles["cover_org"]))
    story.append(_pdf_para(REPORT_TITLE, styles["cover_title"]))
    story.append(_pdf_para(REPORT_SUBTITLE, styles["cover_sub"]))
    story.append(Spacer(1, 4 * cm))
    # 元信息表
    meta_data = [
        [_wrap_long_cell("编 制 人", styles["cover_meta"]), _wrap_long_cell(AUTHOR, styles["cover_meta"])],
        [_wrap_long_cell("编制日期", styles["cover_meta"]), _wrap_long_cell(DATE, styles["cover_meta"])],
    ]
    meta_table = Table(meta_data, colWidths=[4 * cm, 6 * cm])
    meta_table.setStyle(TableStyle([
        ("FONT", (0, 0), (-1, -1), "STSong-Light", 12),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor(GOV_BLUE_HEX)),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("LINEBELOW", (0, 0), (0, -1), 0.5, colors.HexColor(GOV_BLUE_HEX)),
        ("LINEBELOW", (1, 0), (1, -1), 0.5, colors.HexColor(GOV_BLUE_HEX)),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 1 * cm))
    story.append(_pdf_para("—— 内部使用 · 不公开 ——", styles["cover_meta"]))
    story.append(PageBreak())

    # 一、摘要
    story.append(_pdf_para("一、摘要", styles["h1"]))
    story.append(_pdf_para("本报告基于已抓取的 DeepSeek v4-flash 实时价格（api-docs.deepseek.com）以及训练数据估算，"
                           "测算劳动力调查 AI 助手在三档人数（100人/700人/2万人）× 三档月用量（1000/7000/20万次）共 9 档组合下的月度成本。",
                           styles["normal"]))
    story.append(_pdf_para("采购建议档：", styles["no_indent"]))
    story.append(_pdf_para("700 人（市级调查员）× 7000 次/月 ≈ ¥87/月（API ¥24 + 云资源 ¥63）", styles["no_indent"]))
    story.append(_pdf_para("该档覆盖了「市级调查员每月一次集中查询」这一核心真实场景；100 人档可直接降级为本地电脑 + Cloudflare Tunnel，无任何云资源开支。", styles["normal"]))

    story.append(_pdf_para("1.1 关键风险提示", styles["h2"]))
    risks = [
        ("【数据合规】", "本系统架构已严格隔离居民个人信息（H5 → 后端 API → 阿里云百炼 / DeepSeek，仅传查询文本）；"
                          "扩量至市级或省级时，需在采购合同中明确数据流向条款。"),
        ("【政府采购流程】", "云资源月租达一定金额通常需走政府采购或协议供货，域名备案需 15-20 工作日，"
                              "预算应包含备案申请工时。API Key 必须开在单位名下，避免离职交接问题。"),
        ("【量纲校验】", "100 人 × 20 万次/月（每人 67 次/天）属异常档，仅作 API 容量上限测试；"
                          "2 万 × 1000 次/月（每 20 人 1 次/月）使用率过低，建议合并到 7000 次档。"),
    ]
    for title, body in risks:
        p = _pdf_para(f"{title} {body}", styles["normal"])
        story.append(p)

    # 二、9 档对照汇总表
    story.append(_pdf_para("二、9 档对照汇总表", styles["h1"]))
    story.append(_pdf_para("绿色高亮为采购建议档，黄色高亮为异常档（仅作压力测试参考）。", styles["normal"]))
    grid_table = _make_table(GRID_DATA, col_widths=[4.5 * cm, 3.4 * cm, 3.4 * cm, 3.4 * cm], styles=styles)
    extra_cmds = []
    for r_idx, c_idx, fill, font_color in GRID_HIGHLIGHTS:
        extra_cmds.append(("BACKGROUND", (c_idx, r_idx), (c_idx, r_idx), colors.HexColor(fill)))
        extra_cmds.append(("TEXTCOLOR", (c_idx, r_idx), (c_idx, r_idx), colors.HexColor(font_color)))
        extra_cmds.append(("FONT", (c_idx, r_idx), (c_idx, r_idx), "STSong-Light", 11))
    grid_table.setStyle(TableStyle(extra_cmds))
    story.append(grid_table)
    story.append(_pdf_para("采购建议档：700 人 × 7000 次/月 ≈ ¥87/月（API ¥24 + 云资源 ¥63）", styles["note_green"]))
    story.append(_pdf_para("⚠ 异常档：100 人 × 20 万次 = ¥698/月，仅作 API 容量上限参考，不构成实际采购建议。", styles["note_yellow"]))

    story.append(_pdf_para("2.1 价格弹性分析", styles["h2"]))
    bullets = [
        "API 费用（DeepSeek + DashScope）随用量线性增长，几乎纯按量计费。",
        "云资源费用为阶梯式：100 人档 ¥0（本地电脑）；700 人档 ~¥62.5/月（ECS + 域名）；2 万档 ~¥367/月（ECS + SLB + RDS + CDN + 域名）。",
        "采购建议档的 ¥87/月：API ¥24（占 28%）+ 云资源 ¥63（占 72%）。降低 API 用量对总成本影响有限，应聚焦云资源选型。",
    ]
    for b in bullets:
        story.append(_pdf_para("• " + b, styles["bullet"]))

    # 三、详细分档说明
    story.append(_pdf_para("三、详细分档说明", styles["h1"]))

    story.append(_pdf_para("3.1 档位 A：100 人（处室自用）", styles["h2"]))
    story.append(_pdf_para("部署：本地电脑 + Cloudflare Tunnel quick（与现状一致）。", styles["normal"]))
    story.append(_pdf_para("不需要：ECS、域名、备案、CDN。", styles["normal"]))
    story.append(_pdf_para("费用明细（单位：元/月）", styles["h2"]))
    a_data = [
        ["开销项目", "1000 次", "7000 次", "20万次 ⚠"],
        ["LLM 输入（DeepSeek v4-flash）", "¥1.70", "¥11.90", "¥340.00"],
        ["LLM 输出（DeepSeek v4-flash）", "¥0.60", "¥4.20", "¥120.00"],
        ["Embedding（DashScope v3）", "¥1.19", "¥8.33", "¥238.00"],
        ["Cloudflare Tunnel", "免费", "免费", "免费"],
        ["合计", "≈ ¥3/月", "≈ ¥24/月", "≈ ¥698/月"],
    ]
    story.append(_make_table(a_data, col_widths=[6.5 * cm, 3.0 * cm, 3.0 * cm, 3.0 * cm], styles=styles))
    story.append(_pdf_para("⚠ 100 人 × 20 万次 = 每人 2000 次/月（≈67 次/天），属明显异常场景，仅作 API 容量上限参考。",
                           styles["note_yellow"]))
    story.append(_pdf_para("适用判定", styles["h2"]))
    story.append(_pdf_para("• 实际处室用量（每月 100-700 次）：API 月费 < ¥3，可忽略不计。", styles["bullet"]))
    story.append(_pdf_para("• 异常档 ¥698：API 费用已超过部署 ECS 的成本，此规模应迁云。", styles["bullet"]))

    story.append(_pdf_para("3.2 档位 B：700 人（市级使用）", styles["h2"]))
    story.append(_pdf_para("部署：阿里云 ECS（2核4G 突发 t6）+ Cloudflare Tunnel named + 自有域名。", styles["normal"]))
    story.append(_pdf_para("域名备案：政府单位走阿里云 ICP 备案系统（约 15-20 工作日，免费）。", styles["normal"]))
    story.append(_pdf_para("HTTPS：Let's Encrypt 自动续期。", styles["normal"]))
    story.append(_pdf_para("不需要：SLB（单机够用）、CDN（H5 体积小）。", styles["normal"]))
    story.append(_pdf_para("费用明细（单位：元/月）", styles["h2"]))
    b_data = [
        ["开销项目", "1000 次", "7000 次", "20万次"],
        ["LLM 输入", "¥1.70", "¥11.90", "¥340.00"],
        ["LLM 输出", "¥0.60", "¥4.20", "¥120.00"],
        ["Embedding", "¥1.19", "¥8.33", "¥238.00"],
        ["ECS（2核4G t6）", "¥60", "¥60", "¥60"],
        [".cn 域名（年付折月）", "¥2.5", "¥2.5", "¥2.5"],
        ["ICP 备案", "一次性免费", "—", "—"],
        ["Cloudflare Tunnel named", "免费", "免费", "免费"],
        ["SSL 证书", "免费", "免费", "免费"],
        ["合计", "≈ ¥66/月", "≈ ¥87/月", "≈ ¥761/月"],
    ]
    story.append(_make_table(b_data, col_widths=[6.5 * cm, 3.0 * cm, 3.0 * cm, 3.0 * cm], styles=styles))
    story.append(_pdf_para("适用判定", styles["h2"]))
    p = _wrap_long_cell("• 700 × 7000 次/月：月费 ≈ ¥87，<font color='#006100'><b>推荐采购档</b></font>", styles["bullet"])
    story.append(p)
    story.append(_pdf_para("• 700 × 20 万次/月：月费 ≈ ¥761，需要提前扩容 ECS 至 4 核。", styles["bullet"]))
    story.append(_pdf_para("高并发期（7 天承担 60-70% 流量）峰值约 200 次/天（20 万档）。单 ECS 2核4G 跑 DeepSeek 转发绰绰有余，瓶颈在 DeepSeek 端 QPS 限额，不在本机。",
                           styles["normal"]))

    story.append(_pdf_para("3.3 档位 C：2 万人（全省/全市调查员）", styles["h2"]))
    story.append(_pdf_para("部署：阿里云 ECS（4核8G 计算型 c7）+ SLB + Cloudflare Tunnel named + 域名 + 备案。", styles["normal"]))
    story.append(_pdf_para("CDN：H5 静态资源走阿里云 CDN（按流量计费）。", styles["normal"]))
    story.append(_pdf_para("向量库升级：Chroma → 独立服务（Milvus 或 Zvec 备选方案）。", styles["normal"]))
    story.append(_pdf_para("数据库：用户反馈数据上云数据库 RDS（按量计费）。", styles["normal"]))
    story.append(_pdf_para("费用明细（单位：元/月）", styles["h2"]))
    c_data = [
        ["开销项目", "1000 次", "7000 次", "20万次"],
        ["LLM 输入", "¥1.70", "¥11.90", "¥340.00"],
        ["LLM 输出", "¥0.60", "¥4.20", "¥120.00"],
        ["Embedding", "¥1.19", "¥8.33", "¥238.00"],
        ["ECS（4核8G c7）", "¥300", "¥300", "¥300"],
        ["SLB（含流量）", "¥30", "¥30", "¥30"],
        [".cn 域名", "¥2.5", "¥2.5", "¥2.5"],
        ["CDN（H5 静态）", "~¥5", "~¥5", "~¥5"],
        ["RDS MySQL（反馈数据）", "¥30", "¥30", "¥30"],
        ["ICP 备案", "一次性免费", "—", "—"],
        ["Cloudflare Tunnel named", "免费", "免费", "免费"],
        ["合计", "≈ ¥371/月", "≈ ¥392/月", "≈ ¥1066/月"],
    ]
    story.append(_make_table(c_data, col_widths=[6.5 * cm, 3.0 * cm, 3.0 * cm, 3.0 * cm], styles=styles))
    story.append(_pdf_para("适用判定", styles["h2"]))
    story.append(_pdf_para("• 2 万 × 7000 次/月（每人每月 0.35 次 ≈ 每 3 月 1 次）：月费 ≈ ¥392。", styles["bullet"]))
    story.append(_pdf_para("• 2 万 × 20 万次/月（每人每月 10 次 ≈ 每周 2 次）：月费 ≈ ¥1066。", styles["bullet"]))
    story.append(_pdf_para("⚠ 高并发期 7 天承担 14 万次 ≈ 每天 2 万次 ≈ 每分钟 35 次（按 8 小时工作日计）。"
                           "4核8G 单机压力测试需做，极端情况可能需要双 ECS + SLB 横向扩展（月费翻倍至 ¥1500+）。",
                           styles["note_yellow"]))

    # 四、附录 A
    story.append(_pdf_para("四、附录 A：价格核对清单", styles["h1"]))
    story.append(_pdf_para("以下 7 项服务因登录墙 / JS 渲染，自动脚本无法获取精确价格。打印本表后请人工登录核对。", styles["normal"]))
    appendix_a = [
        ["#", "服务", "核对路径", "期望价格", "实际价格（请填）"],
        ["1", "DashScope text-embedding-v3", "https://dashscope.console.aliyun.com/billing", "¥0.7/M tokens", "　　待填"],
        ["2", "阿里云 ECS t6（2核4G）", "https://www.aliyun.com/product/ecs", "~¥60/月", "　　待填"],
        ["3", "阿里云 ECS c7（4核8G）", "同上 → 计算型", "~¥300/月", "　　待填"],
        ["4", "阿里云 SLB", "https://www.aliyun.com/product/slb", "~¥20/月 + 流量", "　　待填"],
        ["5", ".cn 域名", "https://wanwang.aliyun.com", "~¥30/年", "　　待填"],
        ["6", "阿里云 CDN", "https://www.aliyun.com/product/cdn", "~¥0.24/GB", "　　待填"],
        ["7", "阿里云 RDS MySQL", "https://www.aliyun.com/product/rds/mysql", "~¥30/月起", "　　待填"],
    ]
    # Use left-aligned cell style for the path column
    wrapped = []
    for r_idx, row in enumerate(appendix_a):
        new_row = []
        for c_idx, val in enumerate(row):
            if r_idx == 0:
                new_row.append(_wrap_long_cell(val, styles["cell_header"]))
            else:
                if c_idx == 2:
                    new_row.append(_wrap_long_cell(val, styles["cell_left"]))
                else:
                    new_row.append(_wrap_long_cell(val, styles["cell_body"]))
        wrapped.append(new_row)
    a_table = Table(wrapped, colWidths=[0.8 * cm, 3.5 * cm, 5.5 * cm, 2.5 * cm, 2.5 * cm], repeatRows=1)
    a_table.setStyle(_table_style_3line(header_rows=1, body_rows=len(appendix_a) - 1))
    story.append(a_table)
    story.append(_pdf_para("重算公式（每千次）：", styles["no_indent"]))
    story.append(_pdf_para("LLM_每千次 = (1.7 × LLM_输入单价 + 0.3 × LLM_输出单价)", styles["no_indent"]))
    story.append(_pdf_para("Embedding_每千次 = 1.7 × Embedding_单价", styles["no_indent"]))
    story.append(_pdf_para("总月费 = (月次数 ÷ 1000) × (LLM_每千次 + Embedding_每千次) + 云资源固定费", styles["no_indent"]))

    # 五、附录 B
    story.append(_pdf_para("五、附录 B：已确认价格", styles["h1"]))
    story.append(_pdf_para("5.1 DeepSeek v4-flash（自动抓取）", styles["h2"]))
    b2 = [
        ["项目", "美元原价", "换算 CNY（按 USD/CNY ≈ 7.2）"],
        ["输入（cache miss）", "$0.14 / M tokens", "≈ ¥1.0 / M tokens"],
        ["输出", "$0.28 / M tokens", "≈ ¥2.0 / M tokens"],
    ]
    story.append(_make_table(b2, col_widths=[6.0 * cm, 5.0 * cm, 4.5 * cm], styles=styles))
    story.append(_pdf_para("数据来源：https://api-docs.deepseek.com/quick_start/pricing （2026-06-22 抓取）", styles["normal"]))
    story.append(_pdf_para("5.2 Cloudflare Tunnel（公开文档）", styles["h2"]))
    cf = [
        ["模式", "定价", "说明"],
        ["quick tunnel（一次性）", "免费", "无需账号，URL 随机"],
        ["named tunnel（长期）", "免费", "需 Cloudflare 账号 + 域名挂载"],
    ]
    cf_wrapped = []
    for r_idx, row in enumerate(cf):
        new_row = []
        for c_idx, val in enumerate(row):
            if r_idx == 0:
                new_row.append(_wrap_long_cell(val, styles["cell_header"]))
            elif c_idx == 2:
                new_row.append(_wrap_long_cell(val, styles["cell_left"]))
            else:
                new_row.append(_wrap_long_cell(val, styles["cell_body"]))
        cf_wrapped.append(new_row)
    cf_table = Table(cf_wrapped, colWidths=[5.0 * cm, 3.0 * cm, 7.5 * cm], repeatRows=1)
    cf_table.setStyle(_table_style_3line(header_rows=1, body_rows=len(cf) - 1))
    story.append(cf_table)
    story.append(_pdf_para("数据来源：https://developers.cloudflare.com/cloudflare-one/ （2026-06-22 抓取）", styles["normal"]))
    story.append(_pdf_para("Let's Encrypt SSL 证书：免费（自动续期）。", styles["normal"]))

    # 报告结尾
    story.append(Spacer(1, 0.5 * cm))
    p_end = ParagraphStyle("End", parent=styles["normal"], alignment=1, firstLineIndent=0, fontSize=12,
                           textColor=colors.HexColor(GOV_BLUE_HEX))
    story.append(_pdf_para("—— 报告结束 ——", p_end))
    p_sub = ParagraphStyle("Sub", parent=styles["normal"], alignment=1, firstLineIndent=0, fontSize=10,
                           textColor=colors.HexColor(GOV_BLUE_HEX))
    story.append(_pdf_para("本表可作为采购可行性报告附件提交。", p_sub))

    return story


def _make_doc(filename: str = OUTPUT_PDF) -> BaseDocTemplate:
    doc = BaseDocTemplate(
        filename,
        pagesize=A4,
        leftMargin=2.0 * cm,
        rightMargin=2.0 * cm,
        topMargin=2.0 * cm,
        bottomMargin=2.0 * cm,
        title=REPORT_TITLE,
        author=AUTHOR,
    )
    frame = Frame(
        doc.leftMargin,
        doc.bottomMargin,
        doc.width,
        doc.height,
        id="normal",
        showBoundary=0,
    )
    template = PageTemplate(id="default", frames=[frame], onPage=_add_page_decorator)
    doc.addPageTemplates([template])
    return doc


def build_pdf() -> str:
    """Build the PDF using a two-pass strategy so '共 N 页' is accurate.

    Pass 1: build to an in-memory BytesIO to count pages.
    Pass 2: build to file using the known total.
    """
    _register_cn_font()
    styles = _styles()

    from io import BytesIO
    counter_holder = [0]

    def _count_pages(canvas, doc):
        _add_page_decorator(canvas, doc)
        counter_holder[0] = max(counter_holder[0], canvas.getPageNumber())

    # Pass 1: count (in memory)
    mem_path = BytesIO()
    doc_counter = BaseDocTemplate(
        mem_path,
        pagesize=A4,
        leftMargin=2.0 * cm,
        rightMargin=2.0 * cm,
        topMargin=2.0 * cm,
        bottomMargin=2.0 * cm,
        title=REPORT_TITLE,
        author=AUTHOR,
    )
    frame = Frame(doc_counter.leftMargin, doc_counter.bottomMargin, doc_counter.width, doc_counter.height,
                  id="normal", showBoundary=0)
    tpl = PageTemplate(id="default", frames=[frame], onPage=_count_pages)
    doc_counter.addPageTemplates([tpl])
    _TOTAL_PAGES_HOLDER[0] = 0
    story_counter = _build_story(styles)
    doc_counter.build(story_counter)
    total = counter_holder[0]
    if total <= 0:
        total = 1

    # Pass 2: real build to file
    _TOTAL_PAGES_HOLDER[0] = total
    # Write to temp path first, then atomically replace the target.
    # This avoids permission errors if the previous PDF is still open
    # in a viewer on Windows.
    import shutil
    import tempfile
    tmp_fd, tmp_path = tempfile.mkstemp(suffix=".pdf", prefix="cost_report_")
    os.close(tmp_fd)
    doc_real = _make_doc(tmp_path)
    story_real = _build_story(styles)
    doc_real.build(story_real)
    # Replace target (on Windows, this will fail if target is locked).
    try:
        if os.path.exists(OUTPUT_PDF):
            os.remove(OUTPUT_PDF)
        shutil.move(tmp_path, OUTPUT_PDF)
        return OUTPUT_PDF
    except PermissionError:
        # Target is locked by another process (PDF viewer still open).
        # Save under a sibling name with a stamp.
        from datetime import datetime
        stamp = datetime.now().strftime("%H%M%S")
        alt_path = os.path.join(OUTPUT_DIR, f"{OUTPUT_BASENAME}-{stamp}.pdf")
        shutil.move(tmp_path, alt_path)
        return alt_path


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    print("=" * 60)
    print(f"目标目录: {OUTPUT_DIR}")
    print("=" * 60)

    print("\n[1/2] 生成 DOCX ...")
    docx_path = build_docx()
    docx_size = os.path.getsize(docx_path)
    print(f"  OK  -> {docx_path}")
    print(f"       大小: {docx_size:,} bytes ({docx_size / 1024:.1f} KB)")

    print("\n[2/2] 生成 PDF ...")
    pdf_path = build_pdf()
    pdf_size = os.path.getsize(pdf_path)
    print(f"  OK  -> {pdf_path}")
    print(f"       大小: {pdf_size:,} bytes ({pdf_size / 1024:.1f} KB)")

    print("\n" + "=" * 60)
    print("全部完成")
    print("=" * 60)


if __name__ == "__main__":
    main()
