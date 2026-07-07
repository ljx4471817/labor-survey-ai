"""
生成《劳动力调查 AI 助手》向上级汇报的项目介绍 Word 文档
按 GB/T 9704-2012《党政机关公文格式》排版
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------- 字体与排版常量 ----------
TITLE_FONT = '方正小标宋简体'   # 标题
H1_FONT = '黑体'                 # 一级标题
BODY_FONT = '仿宋_GB2312'        # 正文
ASCII_FONT = 'Times New Roman'   # 西文（数字、字母）

TITLE_SIZE = 22   # 二号
H1_SIZE = 16      # 三号
BODY_SIZE = 16    # 三号

LINE_SPACING_PT = 28
INDENT_PT = 32    # 首行缩进 2 字符（三号字 ≈ 16pt × 2）


def set_run_font(run, eastasia, ascii_font=ASCII_FONT, size=BODY_SIZE, bold=False):
    run.font.size = Pt(size)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), eastasia)
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:hAnsi'), ascii_font)


def set_para_spacing(para, line_pt=LINE_SPACING_PT):
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_pt)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, TITLE_FONT, size=TITLE_SIZE)
    set_para_spacing(p)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(20)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, H1_FONT, size=H1_SIZE)
    set_para_spacing(p)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(INDENT_PT)
    run = p.add_run(text)
    set_run_font(run, BODY_FONT)
    set_para_spacing(p)
    return p


def add_signature(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    set_run_font(run, BODY_FONT)
    set_para_spacing(p)
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(0)
    return p


def setup_page(doc):
    section = doc.sections[0]
    section.page_height = Cm(29.7)   # A4
    section.page_width = Cm(21.0)
    # GB/T 9704-2012 标准边距
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)
    section.header_distance = Cm(2.0)
    section.footer_distance = Cm(2.0)


def setup_default_style(doc):
    style = doc.styles['Normal']
    style.font.name = ASCII_FONT
    style.font.size = Pt(BODY_SIZE)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), BODY_FONT)
    rFonts.set(qn('w:ascii'), ASCII_FONT)
    rFonts.set(qn('w:hAnsi'), ASCII_FONT)


def build_document():
    doc = Document()
    setup_page(doc)
    setup_default_style(doc)

    # —— 标题 ——
    add_title(doc, '关于"劳动力调查 AI 助手"项目有关情况的汇报')

    # —— 一、痛点 ——
    add_h1(doc, '一、从两个真实痛点说起')
    add_body(doc,
        '去年下半年起，劳动力调查的填报场景越来越复杂：新型就业形态不断涌现，'
        '"游戏代练""短视频博主""外卖骑手""退休返聘"等身份判定问题，'
        '频繁摆在辅助调查员面前。')
    add_body(doc,
        '【案例 1】游戏代练。某县区辅助调查员入户时，遇到一名 25 岁青年，'
        '全职从事游戏代练，月收入 6,000 余元。调查员拿不准这算不算"就业"，'
        '翻遍制度手册找不到对应条款，第二天上午才敢填。'
        '后来通过电话反复确认，才按"自营劳动者"填报。整个过程耗时近一天。')
    add_body(doc,
        '【案例 2】退休返聘。另一名辅助调查员将一名退休后被原单位返聘的人员'
        '填成了"就业"，区级审核时发现这一判定与制度不符，被打回重做。'
        '被调查户对此也有意见，影响了后续配合度。')
    add_body(doc, '这两个案例的共同点是：标准答案其实就在制度里，但现场查不到、查得慢、查不准。')

    # —— 二、原因 ——
    add_h1(doc, '二、为什么会出现这个问题')
    add_body(doc,
        '一是制度内容多。劳动力调查制度涵盖就业判断、就业身份、单位类型、收入认定等'
        '十余类场景，手册正文 200 余页，辅助调查员无法全部熟记。')
    add_body(doc,
        '二是现场条件受限。入户期间没有电脑、网络、检索工具，'
        '遇到拿不准的指标只能凭经验判断或事后追溯，错填风险较高。')
    add_body(doc,
        '三是经验传承断层。资深调查员积累的"边角案例"判断经验，'
        '往往随人员调动流失，新入职调查员难以快速积累同类经验。')
    add_body(doc, '综上，基层最迫切的需求，是一个能在现场 2 秒钟查到制度依据的工具。')

    # —— 三、方案 ——
    add_h1(doc, '三、我们做的工具：劳动力调查 AI 助手')
    add_body(doc,
        '为回应上述痛点，我们利用业余时间，自主开发了"劳动力调查 AI 助手"，'
        '定位为辅助调查员的随身制度问答工具。它有三个特点：')
    add_body(doc,
        '第一，即开即用。调查员只需在手机上打开一个网页链接（微信内可直接点击），'
        '输入想问的问题，平均 2 秒即可得到答案。无需下载 App，无需注册账号。')
    add_body(doc,
        '第二，答案可追溯。每条回答都附带具体的制度出处——'
        '出自《劳动力调查制度》第几部分、对应哪一条指标，'
        '调查员可据此复核、也可作为审核依据。'
        '系统内部已建立 302 条结构化制度问答，覆盖就业判断、就业身份、单位类型、'
        '特殊场景（港澳台、外籍、服刑人员、僧道人士等）等核心填报场景。')
    add_body(doc,
        '第三，只查制度、不传居民信息。这是本项目最重要的设计原则：'
        '网页端只接收调查员输入的查询文本，不收集、不上传任何'
        '居民姓名、身份证号、地址、家庭收入等个人信息。'
        '系统访问的仅是与制度问答相关的知识库内容。')

    # —— 四、试运行 ——
    add_h1(doc, '四、试运行情况')
    add_body(doc,
        '项目自 2026 年 4 月启动，目前已完成第一阶段建设，处室内测运行中。'
        '成效主要体现在三方面：')
    add_body(doc,
        '一是覆盖度初具规模。知识库已收录 302 条制度问答，覆盖就业判断、就业身份、'
        '单位类型、特殊场景等核心填报场景，'
        '并对游戏代练、退休返聘、外卖骑手等典型坑题进行了专项补全。')
    add_body(doc,
        '二是回答准确率较高。我们建立了 100 题内部评估题库，'
        '覆盖"知识库内""超出知识库""陷阱题""模糊题"四类场景，'
        '实测通过率 99%（其中 99 题答对）。'
        '未通过的一题已定位原因，正在补充知识库条目。')
    add_body(doc,
        '三是具备反馈闭环。每条回答下方都设有"采纳/不采纳"按钮，'
        '调查员的反馈会回流到知识库维护流程，用于持续优化问答质量。')

    # —— 五、成本与合规 ——
    add_h1(doc, '五、花多少钱、安全怎么保')
    add_body(doc,
        '关于成本：按省级 700 名辅助调查员、日均提问 10 次的规模测算，'
        '每月运行成本约 87 元，全年约 1,044 元。'
        '这一规模可走阿里云政府采购协议供货渠道，采购流程简单、周期短。'
        '如未来推广至全国约 2 万名辅助调查员，月成本约 1,066 元，仍属可控范围。')
    add_body(doc, '关于安全合规，我们重点把好四道关：')
    add_body(doc,
        '一是数据边界关。网页端与居民个人信息完全隔离，'
        '仅传输调查员输入的查询文本，从源头杜绝居民数据外泄。')
    add_body(doc,
        '二是模型自主可控。系统采用国产大型语言模型，中文理解能力强、合规可控；'
        '语义匹配采用阿里云通义系列模型，数据不出境。')
    add_body(doc,
        '三是日志限期管理。所有操作日志保留 30 天后滚动清理，不长期留存。')
    add_body(doc,
        '四是备案合规推进。当前为个人主体试运行阶段，仅限处室内部 10–20 人使用；'
        '计划在 3 个月后启动单位主体备案（贵阳调查队作为主办单位），'
        '备案周期约 15–20 个工作日，备案完成后方可面向更大范围推广。')
    add_body(doc,
        '需要如实说明的是，项目目前由 1 名同志利用业余时间开发，迭代速度受限于人力；'
        '知识库条目距 500 条的远期目标还有约 200 条缺口；'
        '试用阶段访问地址每次重启会变化。这些问题将在单位主体备案完成后逐步解决。')

    # —— 六、申请事项 ——
    add_h1(doc, '六、申请事项')
    add_body(doc, '为推动项目从"试用"走向"可用"，恳请支持以下事项：')
    add_body(doc,
        '一是同意扩大试点范围。在完成单位主体备案后，'
        '将使用范围扩展至贵阳市全体辅助调查员，进一步验证大规模场景下的稳定性。')
    add_body(doc,
        '二是启动单位主体备案工作。以贵阳调查队作为主办单位，'
        '启动网站备案及信息系统定级相关流程，预计 15–20 个工作日完成。')
    add_body(doc,
        '三是推进省级协议供货采购。将项目运行费用纳入明年信息化建设预算，'
        '按协议供货方式采购阿里云资源，年支出约 1,044 元。')
    add_body(doc,
        '四是视情况纳入全省推广规划。在贵阳试点成熟的基础上，'
        '研究向全省辅助调查员推广的可行性。')

    # —— 结语 ——
    add_h1(doc, '七、结语')
    add_body(doc,
        '"劳动力调查 AI 助手"是基层同志为解决实际工作难题而自发开展的微创新，'
        '投入小、风险低、见效快、合规边界清晰。'
        '我们坚信，让制度在指尖可用，是提升源头数据质量最朴素、也最有效的方式。'
        '恳请上级机关给予指导和支持，我们将认真落实各项要求，把项目做实做好。')

    # —— 落款 ——
    add_signature(doc, '贵阳市劳动力调查处')
    add_signature(doc, '2026 年 6 月 23 日')

    return doc


if __name__ == '__main__':
    out_path = Path('reports/project-intro-20260623.docx')
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(out_path)
    print(f'已生成：{out_path}')
    print(f'文件大小：{out_path.stat().st_size:,} bytes')