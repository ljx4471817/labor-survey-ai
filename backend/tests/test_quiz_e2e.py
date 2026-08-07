# -*- coding: utf-8 -*-
"""HTTP 级端到端：导入 docx → 提取 → 要点审核 → 生成 → 题目审核 → 下发 → 答题 → 统计（mock LLM）。

PRD v3 10.2。用 httpx ASGITransport 直接驱动 ASGI 应用（避开 TestClient/httpx 版本冲突），
依赖覆盖走 app.dependency_overrides，数据用临时 DB。
"""
from __future__ import annotations

import asyncio
from datetime import datetime, timedelta, timezone

import httpx
import pytest
from docx import Document

from app.infra.auth import require_admin, require_user
from app.main import app
from app.persistence import quiz_db

UTC8 = timezone(timedelta(hours=8))
BASE = "http://test"


@pytest.fixture
def e2e_env(tmp_path, monkeypatch):
    monkeypatch.setattr(quiz_db, "DB_PATH", tmp_path / "quiz_e2e.db")
    quiz_db.reset_conn()
    app.dependency_overrides[require_user] = lambda: "13800000001"
    app.dependency_overrides[require_admin] = lambda: "13900000001"

    def fake_chat(messages, **kwargs):
        sys_prompt = (messages[0]["content"] or "") if messages else ""
        if "提取" in sys_prompt:
            return '[{"section": "审核要点", "content": "家务劳动者无收入应判为非劳动力", "common_error": "误判为就业", "suggest_quiz": true}]'
        return ('{"question": "家务劳动者无收入应判定为？", '
                '"options": {"A": "就业人口", "B": "失业人口", "C": "非劳动力", "D": "在职未就业"}, '
                '"answer": "C", "explanation": "根据审核要点应判为非劳动力。"}')

    monkeypatch.setattr("app.rag.llm.chat", fake_chat)
    monkeypatch.setattr(
        "app.api.quiz_admin.match_kb",
        lambda content: {"faq_id": "023", "question": "家务劳动者如何判定？", "score": 0.87},
    )
    fake_users = [
        {"phone": "13800000001", "name": "张三", "city": "贵阳市", "county": "南明区", "admin_level": "调查员", "active": 1},
        {"phone": "13800000002", "name": "李四", "city": "贵阳市", "county": "云岩区", "admin_level": "调查员", "active": 1},
        {"phone": "13900000001", "name": "管理员", "city": "贵阳市", "county": "", "admin_level": "市级", "active": 1},
    ]
    monkeypatch.setattr("app.api.quiz_admin.whitelist_db.list_all", lambda active_only=False: fake_users)
    monkeypatch.setattr(
        "app.api.quiz_admin.whitelist_db.get_user",
        lambda p: next((u for u in fake_users if u["phone"] == p), None),
    )
    yield
    app.dependency_overrides.clear()


def _make_docx(path) -> None:
    doc = Document()
    doc.add_paragraph("8月劳动力调查工作提示")
    doc.add_paragraph("一、审核要点")
    doc.add_paragraph("将家务劳动者误判为就业是本月常见错误。")
    doc.add_paragraph("二、问卷要点")
    doc.add_paragraph("调查参考周为8月3-9日。")
    doc.save(path)


def test_full_quiz_flow_e2e(e2e_env, tmp_path):
    docx_path = tmp_path / "notice.docx"
    _make_docx(str(docx_path))

    async def run():
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(transport=transport, base_url=BASE, timeout=30) as c:
            # 1) 导入
            with open(docx_path, "rb") as f:
                r = await c.post(
                    "/api/admin/quiz/import",
                    data={"month": "2026-08"},
                    files={"file": ("8月工作提示.docx", f, "application/octet-stream")},
                )
            assert r.status_code == 200, r.text
            quiz_id = r.json()["quiz_id"]

            # 2) 提取（轮询）
            r = await c.post("/api/admin/quiz/extract", json={"quiz_id": quiz_id})
            assert r.status_code == 200, r.text
            task = await _poll(c, "/api/admin/quiz/extract/status/" + r.json()["task_id"])
            assert task["status"] == "done"
            assert task["result"]["keypoints"] >= 1
            assert task["result"]["matched"] >= 1

            # 3) 要点审核
            kps = (await c.get("/api/admin/quiz/keypoints?quiz_id=" + quiz_id)).json()["items"]
            assert kps
            kp = next(k for k in kps if k["kb_match_status"] == "matched")
            r = await c.post("/api/admin/quiz/keypoint/review", json={"keypoint_id": kp["id"], "action": "approve"})
            assert r.status_code == 200

            # 4) 生成（轮询）
            r = await c.post("/api/admin/quiz/generate", json={"quiz_id": quiz_id, "keypoint_ids": [kp["id"]]})
            assert r.status_code == 200, r.text
            task = await _poll(c, "/api/admin/quiz/generate/status/" + r.json()["task_id"])
            assert task["status"] == "done"
            assert task["result"]["questions"] == 1

            # 5) 题目审核
            qs = (await c.get("/api/admin/quiz/questions?quiz_id=" + quiz_id)).json()["items"]
            assert len(qs) == 1
            r = await c.post("/api/admin/quiz/question/review", json={"question_id": qs[0]["id"], "action": "approve"})
            assert r.status_code == 200

            # 6) 下发
            valid_until = (datetime.now(UTC8) + timedelta(days=7)).date().isoformat()
            r = await c.post("/api/admin/quiz/publish", json={
                "quiz_id": quiz_id, "valid_until": valid_until, "targets": ["13800000001"], "action": "publish",
            })
            assert r.status_code == 200, r.text

            # 7) current：未答不泄露答案
            item = (await c.get("/api/quiz/current")).json()["items"][0]
            assert item["answered"] == 0
            qv = item["questions"][0]
            assert "answer" not in qv

            # 8) 提交 → 反馈
            r = await c.post("/api/quiz/submit", json={"quiz_id": quiz_id, "q_id": qv["id"], "selected": "C"})
            assert r.status_code == 200, r.text
            data = r.json()
            assert data["correct"] is True and data["completed"] is True

            # 9) 重复提交 → 409
            r = await c.post("/api/quiz/submit", json={"quiz_id": quiz_id, "q_id": qv["id"], "selected": "A"})
            assert r.status_code == 409

            # 10) current：已答带解析与 KB
            item = (await c.get("/api/quiz/current")).json()["items"][0]
            assert item["answered"] == 1 and item["completed"] is True
            assert item["questions"][0]["kb_ref"]["faq_id"] == "023"

            # 11) 历史
            assert (await c.get("/api/quiz/history")).json()["total"] == 1
            detail = (await c.get("/api/quiz/history/" + quiz_id)).json()
            assert detail["questions"][0]["explanation"]

            # 12) 统计
            st = (await c.get("/api/admin/quiz/stats?quiz_id=" + quiz_id)).json()
            assert st["total_users"] == 1 and st["completed"] == 1 and st["completion_rate"] == 1.0

            # 13) KB 详情
            r = await c.get("/api/faq/023")
            assert r.status_code == 200 and r.json()["id"] == "023"

            # 14) 追加 / 移除未答用户
            r = await c.post("/api/admin/quiz/publish", json={"quiz_id": quiz_id, "targets": ["13800000002"], "action": "append"})
            assert r.status_code == 200 and r.json()["added"] == 1
            r = await c.post("/api/admin/quiz/publish", json={"quiz_id": quiz_id, "targets": ["13800000002"], "action": "remove"})
            assert r.status_code == 200 and r.json()["removed"] == ["13800000002"]

    asyncio.run(run())


async def _poll(client, path, timeout=15):
    import time as _time

    deadline = _time.time() + timeout
    while _time.time() < deadline:
        r = await client.get(path)
        assert r.status_code == 200, r.text
        data = r.json()
        if data["status"] in ("done", "error"):
            return data
        await asyncio.sleep(0.3)
    raise AssertionError(f"任务超时: {path}")
